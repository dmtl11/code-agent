from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


SOURCE_PATH = Path("docs/Code_Agent_实现与扩展功能说明.docx")
OUTPUT_PATH = Path("docs/Code_Agent_实现与扩展功能说明_含监控实现.docx")


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(text, style="Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_code(document: Document, text: str):
    paragraph = document.add_paragraph(text, style="Code Block")
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def main() -> None:
    document = Document(SOURCE_PATH)
    document.add_page_break()

    document.add_heading("八、面向 DeepSeek Harness 的可观测性与项目监控", level=1)
    add_body(
        document,
        "在前面的 ReAct、Repo Map、局部编辑、上下文管理和 Session 持久化基础上，本次继续补充项目监控能力。目标不是复制某个产品的内部实现，而是借鉴 DeepSeek Harness 强调的“每次运行都可追踪”思想：一次用户请求从模型调用、工具调用、上下文压缩到最终结果，都能够留下可聚合的运行数据。"
    )

    document.add_heading("8.1 设计依据与实现边界", level=2)
    add_body(
        document,
        "DeepSeek Harness 的公开介绍强调 append-only 的 Session Log 和完整的 trajectory view，会记录系统提示词、推理过程、工具调用与结果、子智能体调度以及上下文注入。本项目采用更轻量的本地实现：详细过程继续保存在 Session 的 events 表中，适合统计的指标另外写入 SQLite 的 metric_records 表。这样既能回放一次任务，也能按工作区、Session、模型和工具进行汇总。"
    )
    add_body(
        document,
        "DeepSeek API 返回的 usage 字段包含 prompt_tokens、completion_tokens、total_tokens，以及 prompt cache hit/miss 等字段。因此模型适配层先把不同供应商的 usage 统一成内部格式，再交给 Agent 和 SessionStore 记录；没有返回 usage 的兼容模型只做近似估算，并显式标注为 estimated，不把估算值当作官方计费数据。"
    )
    add_bullet(document, "已实现：模型请求量、Token 用量、缓存命中、延迟、运行错误率、工具错误率、上下文压缩和最近错误。")
    add_bullet(document, "已实现：Workspace 级和 Session 级聚合，可按模型和工具查看明细。")
    add_bullet(document, "当前边界：数据存储在本地 SQLite，尚未接入 Prometheus、OpenTelemetry、告警系统或费用账单。")
    add_body(
        document,
        "参考资料：DeepSeek Harness 介绍页：https://www.deepseek.com/harness/en/；DeepSeek Chat Completions API：https://api-docs.deepseek.com/api/create-chat-completion/；Token usage 说明：https://api-docs.deepseek.com/quick_start/token_usage/。"
    )

    document.add_heading("8.2 query 到监控数据的完整路径", level=2)
    add_body(
        document,
        "用户在网页输入 query 后，流程与普通 Code 模式相同，但每个关键步骤都额外生成一个结构化事件。请求先进入 server.py 的流式执行入口，Agent 根据当前模式、Repo Map 和 Session 历史组装 messages，然后调用统一的 ChatModel。模型返回工具调用时，LocalTools 执行读文件、搜索、局部替换或运行命令，并把工具结果回传给模型。"
    )
    add_code(
        document,
        "用户 query\n  -> server.py 创建 run_id\n  -> agent.py 组装 Repo Map、历史和当前任务\n  -> model.py 调用选中的 LLM\n  -> emit(llm_call / llm_error)\n  -> LocalTools 执行工具\n  -> emit(tool_result)\n  -> ContextManager 发生压缩时 emit(context)\n  -> SessionStore.metric_records 持久化\n  -> /api/monitoring 聚合\n  -> Web Monitor 面板展示"
    )
    add_body(
        document,
        "server.py 为一次运行创建唯一 run_id，并在 emit 函数中把 run_id 注入每条事件。模型调用、工具结果和上下文事件一边通过 NDJSON 推送到网页，一边写入 SessionStore。Agent 结束后再写入一条 run 记录，用于计算整次任务的完成率和运行延迟。"
    )
    add_code(
        document,
        "def emit(event):\n    event = {**event, \"run_id\": run_id}\n    store.append_event(session_id, event)\n    if event[\"type\"] in {\"llm_call\", \"llm_error\",\n                            \"tool_result\", \"context\"}:\n        store.record_metric(session_id, event)\n    stream.write(event)"
    )

    document.add_heading("8.3 模型 usage 的统一与真实度标记", level=2)
    add_body(
        document,
        "不同模型供应商的字段名并不完全一致：OpenAI-compatible 接口通常使用 prompt_tokens 和 completion_tokens，Anthropic 使用 input_tokens 和 output_tokens，DeepSeek 还可能提供缓存命中和未命中 Token。model.py 的 _normalize_usage 将它们转换成 prompt_tokens、completion_tokens、total_tokens、prompt_cache_hit_tokens 和 prompt_cache_miss_tokens。"
    )
    add_code(
        document,
        "def _normalize_usage(usage):\n    prompt = usage.get(\"prompt_tokens\", usage.get(\"input_tokens\", 0))\n    completion = usage.get(\"completion_tokens\", usage.get(\"output_tokens\", 0))\n    total = usage.get(\"total_tokens\", prompt + completion)\n    return {\n        \"prompt_tokens\": prompt,\n        \"completion_tokens\": completion,\n        \"total_tokens\": total,\n        \"prompt_cache_hit_tokens\": usage.get(\"prompt_cache_hit_tokens\", 0),\n        \"prompt_cache_miss_tokens\": usage.get(\"prompt_cache_miss_tokens\", 0),\n    }"
    )
    add_body(
        document,
        "如果接口返回 usage，ChatModel 将 usage_source 设置为 actual；如果供应商没有返回 usage，Agent 使用 ContextManager 的估算器计算输入和输出 Token，并设置为 estimated。网页将两者分开显示，使用户能够判断监控数字的可信度。"
    )

    document.add_heading("8.4 metric_records 的结构与写入时机", level=2)
    add_body(
        document,
        "metric_records 是面向统计的窄表，每行表示一个可观测动作，而不是保存完整对话内容。核心字段包括 session_id、run_id、kind、provider、model、tool_name、ok、latency_ms、Token 字段、usage_source、error、compacted_blocks 和 truncated_tool_results。详细的参数和输出仍在 Session events 中保存，避免统计表过度膨胀。"
    )
    add_code(
        document,
        "CREATE TABLE IF NOT EXISTS metric_records (\n    id INTEGER PRIMARY KEY AUTOINCREMENT,\n    session_id TEXT NOT NULL,\n    run_id TEXT,\n    kind TEXT NOT NULL,\n    provider TEXT, model TEXT, tool_name TEXT,\n    ok INTEGER NOT NULL, latency_ms REAL,\n    prompt_tokens INTEGER, completion_tokens INTEGER,\n    total_tokens INTEGER, usage_source TEXT, error TEXT,\n    compacted_blocks INTEGER, truncated_tool_results INTEGER,\n    created_at TEXT NOT NULL\n)"
    )
    add_body(
        document,
        "模型调用成功写入 llm_call，模型异常写入 llm_error，工具执行完成写入 tool_result，上下文管理器发生压缩或截断时写入 context，整个 Agent 结束时写入 run。clear_session 会同时清理这些指标，避免删除 Session 后留下孤立数据。"
    )

    document.add_heading("8.5 指标口径", level=2)
    add_body(document, "monitoring_summary 在 SessionStore 中完成聚合，当前主要指标口径如下：")
    add_bullet(document, "LLM error rate = llm_error 数 / (llm_call 数 + llm_error 数) × 100%。")
    add_bullet(document, "Run error rate = 失败 run 数 / 全部 run 数 × 100%，反映一次用户任务是否完成。")
    add_bullet(document, "Tool error rate = ok=false 的 tool_result 数 / 全部工具调用数 × 100%。")
    add_bullet(document, "平均延迟是所有已记录动作 latency_ms 的平均值，P95 使用排序后的 95 分位值。")
    add_bullet(document, "Prompt cache hit rate = cache hit Token / (cache hit Token + cache miss Token) × 100%。")
    add_bullet(document, "actual_usage_calls 和 estimated_usage_calls 分开统计，避免把估算 Token 混入真实用量判断。")
    add_body(
        document,
        "当分母为 0 时，错误率和缓存命中率返回 0.0，而不是返回 NaN 或空值，前端可以直接展示。没有执行过任务的 Workspace 因此会显示全 0，这是初始化状态，不代表模型调用失败。"
    )

    document.add_heading("8.6 API 与网页 Monitor 面板", level=2)
    add_body(
        document,
        "server.py 新增 GET /api/monitoring 接口。默认参数按 workspace 聚合；追加 session_id 后只返回当前 Session 的数据。接口返回 summary、providers、tools、recent_errors 和 last_updated，前端不需要理解 SQLite 结构，只负责展示聚合结果。"
    )
    add_code(
        document,
        "GET /api/monitoring?workspace=workspace\nGET /api/monitoring?workspace=workspace&session_id=<session-id>\n\n{\n  \"summary\": {\n    \"llm_requests\": 12,\n    \"llm_error_rate\": 8.33,\n    \"total_tokens\": 18420,\n    \"avg_latency_ms\": 920,\n    \"p95_latency_ms\": 2100,\n    \"compactions\": 2\n  },\n  \"providers\": [],\n  \"tools\": [],\n  \"recent_errors\": []\n}"
    )
    add_body(
        document,
        "右侧 Agent 区域新增 Chat、Review、Monitor 三个标签。Monitor 中展示 LLM 请求量、LLM/Run/Tool 错误率、总 Token 和实际 usage 次数、平均/P95 延迟、Prompt 缓存命中率、工具调用及上下文压缩次数，并按模型、工具和最近错误分组。点击 Refresh 会重新请求接口；运行任务完成后也会自动刷新。"
    )

    document.add_heading("8.7 如何验证监控能力确实生效", level=2)
    add_body(document, "验证分为自动化测试和手动观察两部分。自动化测试覆盖以下场景：")
    add_bullet(document, "Fake model 返回 usage 时，验证 OpenAI-compatible 和 Anthropic 字段都能归一化。")
    add_bullet(document, "写入一次成功模型调用、一次失败模型调用、一次工具调用和一次 context 事件，验证汇总出的请求数、错误率、Token 和压缩次数。")
    add_bullet(document, "通过 HTTP 请求访问 /api/monitoring，验证 Workspace 聚合结果能够返回给前端。")
    add_bullet(document, "运行 Python、C++ 文件或执行一个 Agent 任务后，在 Monitor 标签中观察请求、延迟、工具和错误记录。")
    add_code(
        document,
        "python -m unittest discover -s tests -p \"test_*.py\"\n# 预期：18/18 tests passed\n\n# 浏览器打开\nhttp://127.0.0.1:8767/\n# 进入右侧 Monitor 标签，点击 Refresh"
    )

    document.add_heading("8.8 与 DeepSeek Harness 的接近点和后续方向", level=2)
    add_body(
        document,
        "本次实现已经在本地 Harness 中建立了可观测性的最小闭环：每次运行有 run_id，每个动作有结构化记录，Token 和错误可以聚合，网页可以按 Workspace 或 Session 查看。它接近的是 DeepSeek Harness 的 traceability 思路，而不是宣称复刻其完整内部平台。"
    )
    add_bullet(document, "下一步可把 metric_records 导出为 OpenTelemetry spans，接入 Jaeger、Grafana 或其他 OTLP 后端。")
    add_bullet(document, "下一步可增加按时间窗口的趋势序列、模型成本估算、阈值告警和失败任务重试统计。")
    add_bullet(document, "下一步可为事件增加 trace_id、parent_span_id 和 tool_call_id，形成更接近分布式追踪的树形视图。")
    add_bullet(document, "下一步可把长时间运行的任务、并发 Session 和取消操作纳入监控，补充活跃任务数与取消率。")

    document.core_properties.comments = "补充面向 DeepSeek Harness 可观测性思路的监控实现说明"
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
